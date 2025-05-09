import os
import requests
from bs4 import BeautifulSoup
import urllib.parse
import pytesseract
from PIL import Image
import re
import spacy
import chardet
from typing import List, Optional
import tempfile
import logging

# PDF text extraction
import PyPDF2

# Additional document handling
import docx
import csv
import openpyxl
from datetime import datetime

class WebScraper:
    def __init__(self, 
                 url: str, 
                 base_output_dir: str = tempfile.gettempdir(),
                 extract_text: bool = False,
                 extract_links: bool = False,
                 extract_documents: bool = False,
                 extract_images: bool = False,
                 flat_structure: bool = False):
        """
        Initialize the web scraper with precise extraction control
        
        :param url: Target webpage URL
        :param base_output_dir: Base directory for saving scraped content
        :param extract_text: Flag to extract webpage text
        :param extract_links: Flag to extract hyperlinks
        :param extract_documents: Flag to download documents
        :param extract_images: Flag to download images
        :param flat_structure: Flag to use flat directory structure (no subdirectories)
        """
        # Configure logging
        logging.basicConfig(
            level=logging.INFO, 
            format='%(asctime)s - %(levelname)s: %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Validate URL
        if not url or not isinstance(url, str):
            raise ValueError("URL must be a valid string")
        
        if not url.startswith(('http://', 'https://')):
            raise ValueError("URL must start with http:// or https://")

        # Core configuration
        self.url = url
        self.BeautifulSoup = BeautifulSoup
        
        # Strict extraction flags
        self.extract_text = extract_text
        self.extract_link = extract_links
        self.extract_documents = extract_documents
        self.extract_images = extract_images
        self.flat_structure = flat_structure
        
        # Validate base_output_dir
        if not os.path.exists(base_output_dir):
            try:
                os.makedirs(base_output_dir, exist_ok=True)
                self.logger.info(f"Created base output directory: {base_output_dir}")
            except Exception as e:
                raise IOError(f"Failed to create base output directory: {str(e)}")
        
        # Create safe, unique folder name
        self.safe_folder_name = self._create_safe_folder_name(url)
        
        # Create output paths
        self.output_dir = os.path.join(base_output_dir, self.safe_folder_name)
        
        # Set up directory structure based on flat_structure setting
        if self.flat_structure:
            self.links_dir = self.output_dir
            self.docs_dir = self.output_dir
            self.images_dir = self.output_dir
        else:
            self.links_dir = os.path.join(self.output_dir, 'links')
            self.docs_dir = os.path.join(self.output_dir, 'documents')
            self.images_dir = os.path.join(self.output_dir, 'images')
        
        # Conditionally create directories only if extraction is enabled
        try:
            if self.extract_text or self.extract_link or self.extract_documents or self.extract_images:
                os.makedirs(self.output_dir, exist_ok=True)
                
            if not self.flat_structure:
                if self.extract_link:
                    os.makedirs(self.links_dir, exist_ok=True)
                if self.extract_documents:
                    os.makedirs(self.docs_dir, exist_ok=True)
                if self.extract_images:
                    os.makedirs(self.images_dir, exist_ok=True)
        except Exception as e:
            raise IOError(f"Failed to create output directories: {str(e)}")
            
        # Check dependencies if extraction is enabled
        if extract_images:
            self._check_pytesseract()
            
        if extract_text or extract_documents or extract_images:
            self._check_spacy()
        
        # Load spaCy model for text processing if needed
        self.nlp = None
        if extract_text or extract_documents or extract_images:
            self._load_spacy_model()

    def _check_pytesseract(self):
        """
        Check if pytesseract is properly installed and configured
        """
        try:
            import pytesseract
            # Test pytesseract with a small operation
            version = pytesseract.get_tesseract_version()
            self.logger.info(f"Tesseract OCR version: {version}")
        except ImportError:
            self.logger.error("pytesseract module is not installed. OCR functionality will be disabled.")
            self.logger.error("Install using: pip install pytesseract")
            self.extract_images = False
        except (pytesseract.TesseractNotFoundError, Exception) as e:
            self.logger.error(f"Tesseract OCR is not properly installed: {str(e)}")
            self.logger.error("Please install Tesseract OCR on your system: https://github.com/tesseract-ocr/tesseract")
            self.extract_images = False
            
    def _check_spacy(self):
        """
        Check if spaCy is properly installed
        """
        try:
            import spacy
            self.logger.info("spaCy is installed and available.")
        except ImportError:
            self.logger.error("spaCy module is not installed. Text processing functionality will be limited.")
            self.logger.error("Install using: pip install spacy")

    def _load_spacy_model(self):
        """
        Load or download spaCy English model with improved error handling
        """
        if self.nlp is not None:
            return  # Model already loaded
            
        try:
            import spacy
            self.nlp = spacy.load('en_core_web_sm')
            self.logger.info("SpaCy model loaded successfully")
        except OSError:
            # Model not found, try to download
            try:
                self.logger.info("Downloading spaCy English model...")
                from spacy.cli import download
                download('en_core_web_sm')
                self.nlp = spacy.load('en_core_web_sm')
                self.logger.info("SpaCy model downloaded and loaded successfully")
            except Exception as e:
                self.logger.error(f"Failed to download spaCy model: {str(e)}")
                self.logger.warning("Using basic text cleaning without spaCy")
                # Create a minimal fallback for text cleaning
                self.nlp = None
        except ImportError:
            self.logger.error("spaCy module not found. Install using: pip install spacy")
            self.nlp = None
        except Exception as e:
            self.logger.error(f"Unexpected error loading spaCy model: {str(e)}")
            self.nlp = None

    def _create_safe_folder_name(self, url: str) -> str:
        """
        Create a safe, unique folder name from URL
        """
        safe_name = url.replace('https://', '').replace('http://', '')
        safe_name = "".join(c for c in safe_name if c.isalnum() or c in ['-', '_'])
        
        # Add timestamp (using ISO format for readability and filesystem compatibility)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Combine safe name and timestamp
        full_safe_name = f"{safe_name[:20]}_{timestamp}"
        return full_safe_name

    def fetch_page_content(self):
        """
        Fetch webpage content with robust error handling
        """
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(self.url, headers=headers, timeout=10)
            response.raise_for_status()
            return response
        except requests.RequestException as e:
            self.logger.error(f"Webpage fetch failed: {e}")
            return None
        
    def detect_encoding(self, filepath: str) -> str:
        """
        Detect the encoding of a file
        
        :param filepath: Path to the file
        :return: Detected encoding
        """
        with open(filepath, 'rb') as file:
            result = chardet.detect(file.read())
        return result['encoding'] or 'utf-8'

    def extract_links(self, soup: BeautifulSoup) -> List[str]:
        """
        Selectively extract links based on configuration
        """
        if not self.extract_link:
            return []

        links = set()
        base_url = urllib.parse.urlparse(self.url).scheme + "://" + urllib.parse.urlparse(self.url).netloc
        
        for a_tag in soup.find_all('a', href=True):
            href = a_tag['href']
            full_url = urllib.parse.urljoin(base_url, href)
            links.add(full_url)
        
        return list(links)

    def save_links(self, links: List[str]):
        """
        Save extracted links with strict extraction control
        """
        if not self.extract_link:
            return 0
        saved_link = 0
        
        # Use prefix for flat structure
        if self.flat_structure:
            links_file = os.path.join(self.links_dir, 'links_extracted.txt')
        else:
            links_file = os.path.join(self.links_dir, 'links.txt')
            
        try:
            with open(links_file, 'w', encoding='utf-8') as f:
                for link in links:
                    f.write(f"{link}\n")
                    saved_link += 1
            self.logger.info(f"Saved {len(links)} links")
        except Exception as e:
            self.logger.error(f"Link saving failed: {e}")
        return saved_link

    def download_documents(self, soup: BeautifulSoup):
        """
        Strictly download documents only if enabled with enhanced error handling
        """
        if not self.extract_documents:
            return 0

        doc_extensions = ['.pdf', '.docx', '.doc', '.txt', '.csv', '.xls', '.xlsx']
        base_url = urllib.parse.urlparse(self.url).scheme + "://" + urllib.parse.urlparse(self.url).netloc
        
        documents_downloaded = 0
        documents_skipped = 0
        documents_failed = 0
        
        # Find all links
        all_links = soup.find_all('a', href=True)
        doc_links = []
        
        # Filter for document links
        for a_tag in all_links:
            href = a_tag['href']
            if any(ext in href.lower() for ext in doc_extensions):
                doc_links.append(a_tag)
                
        if len(doc_links) == 0:
            self.logger.info("No document links found on the page")
            return 0
            
        self.logger.info(f"Found {len(doc_links)} potential document links")
        
        for a_tag in doc_links:
            href = a_tag['href']
            full_url = urllib.parse.urljoin(base_url, href)
            
            try:
                # Validate URL before downloading
                if not self._validate_url(full_url):
                    self.logger.warning(f"Skipping document download from invalid URL: {full_url}")
                    documents_skipped += 1
                    continue
                    
                # Log download attempt
                self.logger.info(f"Downloading document: {full_url}")
                
                # Download with timeout and headers
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                }
                doc_response = requests.get(full_url, headers=headers, timeout=15)
                doc_response.raise_for_status()
                
                # Check content type
                content_type = doc_response.headers.get('Content-Type', '').lower()
                if 'html' in content_type and not any(ext in href.lower() for ext in ['.htm', '.html']):
                    self.logger.warning(f"Content appears to be HTML, not a document: {full_url}")
                    documents_skipped += 1
                    continue
                
                # Add prefix for flat structure
                if self.flat_structure:
                    filename = os.path.join(self.docs_dir, f"doc_{os.path.basename(href)}")
                else:
                    filename = os.path.join(self.docs_dir, os.path.basename(href))
                
                # Ensure the filename is safe and unique
                safe_filename = self._create_safe_filename(filename)
                
                # Write to file
                with open(safe_filename, 'wb') as f:
                    f.write(doc_response.content)
                
                documents_downloaded += 1
                self.logger.info(f"Successfully downloaded document: {safe_filename}")
                
                # Extract text from document
                self.extract_text_from_document(safe_filename)
                
            except requests.RequestException as e:
                self.logger.warning(f"Document download failed: {str(e)}")
                documents_failed += 1
            except Exception as e:
                self.logger.warning(f"Document processing failed: {str(e)}")
                documents_failed += 1
        
        self.logger.info(f"Documents: {documents_downloaded} downloaded, {documents_skipped} skipped, {documents_failed} failed")
        return documents_downloaded
        
    def _create_safe_filename(self, original_filename):
        """
        Create a safe filename by removing invalid characters and ensuring uniqueness
        
        :param original_filename: Original filename path
        :return: Safe filename path
        """
        directory = os.path.dirname(original_filename)
        filename = os.path.basename(original_filename)
        
        # Remove invalid characters
        safe_name = re.sub(r'[\\/*?:"<>|]', '_', filename)
        
        # Ensure the name isn't too long
        if len(safe_name) > 100:
            name, ext = os.path.splitext(safe_name)
            safe_name = name[:95] + ext
            
        # Make sure we have a unique filename
        counter = 1
        final_path = os.path.join(directory, safe_name)
        
        while os.path.exists(final_path):
            name, ext = os.path.splitext(safe_name)
            safe_name = f"{name}_{counter}{ext}"
            final_path = os.path.join(directory, safe_name)
            counter += 1
            
        return final_path

    def extract_text_from_document(self, filepath: str):
        """
        Extract text from various document types with robust error handling
        
        :param filepath: Path to the document file
        """
        # If document extraction is disabled, skip
        if not self.extract_documents:
            return

        try:
            # Check if file exists and is accessible
            if not os.path.exists(filepath):
                self.logger.error(f"Document file not found: {filepath}")
                return
                
            # Check file size to avoid processing very large documents
            file_size_mb = os.path.getsize(filepath) / (1024 * 1024)
            if file_size_mb > 50:  # 50MB limit
                self.logger.warning(f"Document file too large ({file_size_mb:.2f} MB), extraction may be slow: {filepath}")
            
            # Get file extension
            file_ext = os.path.splitext(filepath)[1].lower()
            
            text = ""
            extraction_success = False

            # PDF extraction
            if file_ext == '.pdf':
                try:
                    text = self.extract_pdf_text(filepath)
                    extraction_success = True
                except ImportError as e:
                    self.logger.error(f"PDF extraction requires PyPDF2: {str(e)}")
                except Exception as e:
                    self.logger.error(f"Error extracting text from PDF {filepath}: {str(e)}")
            
            # DOCX extraction
            elif file_ext in ['.docx', '.doc']:
                try:
                    text = self.extract_docx_text(filepath)
                    extraction_success = True
                except ImportError as e:
                    self.logger.error(f"DOCX extraction requires python-docx: {str(e)}")
                except Exception as e:
                    self.logger.error(f"Error extracting text from DOCX {filepath}: {str(e)}")
            
            # CSV extraction
            elif file_ext == '.csv':
                try:
                    text = self.extract_csv_text(filepath)
                    extraction_success = True
                except Exception as e:
                    self.logger.error(f"Error extracting text from CSV {filepath}: {str(e)}")
            
            # Excel extraction
            elif file_ext in ['.xls', '.xlsx']:
                try:
                    text = self.extract_excel_text(filepath)
                    extraction_success = True
                except ImportError as e:
                    self.logger.error(f"Excel extraction requires openpyxl: {str(e)}")
                except Exception as e:
                    self.logger.error(f"Error extracting text from Excel {filepath}: {str(e)}")
            
            # Plain text
            elif file_ext == '.txt':
                try:
                    # Detect encoding and read file
                    encoding = self.detect_encoding(filepath)
                    with open(filepath, 'r', encoding=encoding) as f:
                        text = f.read()
                    extraction_success = True
                except Exception as e:
                    self.logger.error(f"Error reading text file {filepath}: {str(e)}")
            
            else:
                self.logger.warning(f"Unsupported file type: {file_ext} - {filepath}")
                return

            # If extraction failed, stop here
            if not extraction_success or not text:
                self.logger.warning(f"No text extracted from document: {filepath}")
                return
                
            # Clean text
            cleaned_text = self.clean_text(text)
            
            # Skip if no meaningful text was extracted
            if not cleaned_text or len(cleaned_text.strip()) < 10:
                self.logger.info(f"No meaningful text extracted from document: {filepath}")
                return
                
            # Add prefix for flat structure
            if self.flat_structure:
                basename = f"doc_{os.path.splitext(os.path.basename(filepath))[0]}.txt"
            else:
                basename = f"{os.path.splitext(os.path.basename(filepath))[0]}.txt"
                
            # Save text to file
            text_filename = os.path.join(self.docs_dir, basename)
            with open(text_filename, 'w', encoding='utf-8') as f:
                f.write(cleaned_text)
                
            self.logger.info(f"Successfully extracted text from document: {filepath}")
        
        except Exception as e:
            self.logger.error(f"Error extracting text from document {filepath}: {str(e)}")

    def extract_pdf_text(self, filepath: str) -> str:
        """
        Extract text from PDF file with error handling
        
        :param filepath: Path to PDF file
        :return: Extracted text
        """
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF extraction. Install using: pip install PyPDF2")
            
        text = ""
        try:
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if reader.is_encrypted:
                    self.logger.warning(f"PDF is encrypted, attempting to decrypt: {filepath}")
                    try:
                        # Try with empty password
                        reader.decrypt('')
                    except:
                        self.logger.error(f"Failed to decrypt PDF: {filepath}")
                        return "PDF is encrypted and could not be decrypted."
                
                # Check page count
                page_count = len(reader.pages)
                if page_count > 100:
                    self.logger.warning(f"PDF has {page_count} pages, extraction may be slow: {filepath}")
                
                # Extract text from each page
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"
                        
                        # Log progress for large documents
                        if page_count > 20 and page_num % 10 == 0:
                            self.logger.info(f"Extracted {page_num}/{page_count} pages from {filepath}")
                            
                    except Exception as e:
                        self.logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                
        except Exception as e:
            self.logger.error(f"Error processing PDF file {filepath}: {str(e)}")
            raise
            
        return text.strip()

    def extract_docx_text(self, filepath: str) -> str:
        """
        Extract text from DOCX file with error handling
        
        :param filepath: Path to DOCX file
        :return: Extracted text
        """
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX extraction. Install using: pip install python-docx")
            
        try:
            doc = docx.Document(filepath)
            text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX file {filepath}: {str(e)}")
            raise

    def extract_csv_text(self, filepath: str) -> str:
        """
        Extract text from CSV file with error handling
        
        :param filepath: Path to CSV file
        :return: Extracted text
        """
        text = []
        
        try:
            # Detect encoding 
            encoding = self.detect_encoding(filepath)
            
            # Try multiple delimiters
            delimiters = [',', ';', '\t', '|']
            best_delimiter = ','
            max_columns = 0
            
            # Find the best delimiter by checking which one gives the most columns
            for delimiter in delimiters:
                try:
                    with open(filepath, 'r', encoding=encoding) as csvfile:
                        sample = csvfile.read(2048)  # Read a sample
                        if delimiter in sample:
                            sample_rows = sample.split('\n')
                            if len(sample_rows) > 1:
                                cols = len(sample_rows[0].split(delimiter))
                                if cols > max_columns:
                                    max_columns = cols
                                    best_delimiter = delimiter
                except Exception:
                    continue
            
            # Read with best delimiter
            with open(filepath, 'r', encoding=encoding) as csvfile:
                import csv
                csv_reader = csv.reader(csvfile, delimiter=best_delimiter)
                for row in csv_reader:
                    text.append(" ".join([str(cell) for cell in row if cell]))
                    
            return "\n".join(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from CSV file {filepath}: {str(e)}")
            raise

    def extract_excel_text(self, filepath: str) -> str:
        """
        Extract text from Excel file with error handling
        
        :param filepath: Path to Excel file
        :return: Extracted text
        """
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for Excel extraction. Install using: pip install openpyxl")
            
        text = []
        
        try:
            workbook = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            
            # Track progress for large files
            sheet_count = len(workbook.sheetnames)
            if sheet_count > 10:
                self.logger.info(f"Processing large Excel file with {sheet_count} sheets: {filepath}")
            
            for sheet_name in workbook.sheetnames:
                try:
                    sheet = workbook[sheet_name]
                    
                    # Add sheet name as header
                    text.append(f"Sheet: {sheet_name}")
                    
                    # Process rows
                    row_count = 0
                    for row in sheet.iter_rows(values_only=True):
                        row_values = [str(cell) for cell in row if cell is not None]
                        if row_values:  # Skip empty rows
                            text.append(" | ".join(row_values))
                        row_count += 1
                        
                        # Log progress for large sheets
                        if row_count > 1000 and row_count % 1000 == 0:
                            self.logger.info(f"Processed {row_count} rows in sheet {sheet_name}")
                            
                    text.append("")  # Add separator between sheets
                except Exception as e:
                    self.logger.warning(f"Error processing sheet {sheet_name}: {str(e)}")
                    continue
            
            return "\n".join(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from Excel file {filepath}: {str(e)}")
            raise

    def download_images(self, soup: BeautifulSoup):
        """
        Strictly download images only if enabled with robust error handling
        """
        if not self.extract_images:
            return 0

        base_url = urllib.parse.urlparse(self.url).scheme + "://" + urllib.parse.urlparse(self.url).netloc
        
        images_downloaded = 0
        images_skipped = 0
        images_failed = 0
        
        # Find all image tags
        all_images = soup.find_all('img', src=True)
        
        if len(all_images) == 0:
            self.logger.info("No images found on the page")
            return 0
            
        self.logger.info(f"Found {len(all_images)} potential images")
        
        # Try additional image sources (background images, etc.)
        for elem in soup.find_all(style=True):
            style = elem['style']
            url_matches = re.findall(r'url\([\'"]?(.*?)[\'"]?\)', style)
            
            for url_match in url_matches:
                if url_match and not url_match.startswith('data:'):  # Skip data URLs
                    img_tag = soup.new_tag('img')
                    img_tag['src'] = url_match
                    all_images.append(img_tag)
        
        for img_tag in all_images:
            img_url = img_tag['src']
            
            # Skip data URLs
            if img_url.startswith('data:'):
                self.logger.debug("Skipping data URL")
                images_skipped += 1
                continue
                
            # Skip small icons and spacer images
            if 'icon' in img_url.lower() or 'spacer' in img_url.lower():
                if img_tag.get('width') and int(img_tag['width']) < 50:
                    self.logger.debug(f"Skipping small icon: {img_url}")
                    images_skipped += 1
                    continue
            
            # Build full URL
            full_img_url = urllib.parse.urljoin(base_url, img_url)
            
            try:
                # Validate URL before downloading
                if not self._validate_url(full_img_url):
                    self.logger.warning(f"Skipping image download from invalid URL: {full_img_url}")
                    images_skipped += 1
                    continue
                    
                # Log download attempt
                self.logger.info(f"Downloading image: {full_img_url}")
                
                # Download with timeout and headers
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    'Referer': self.url  # Some servers require referer
                }
                img_response = requests.get(full_img_url, headers=headers, timeout=10)
                img_response.raise_for_status()
                
                # Check content type
                content_type = img_response.headers.get('Content-Type', '').lower()
                if not content_type.startswith('image/'):
                    self.logger.warning(f"Content is not an image ({content_type}): {full_img_url}")
                    images_skipped += 1
                    continue
                    
                # Check file size
                if len(img_response.content) < 100:  # Extremely small files are usually not valid images
                    self.logger.warning(f"Image file too small, likely not valid: {full_img_url}")
                    images_skipped += 1
                    continue
                
                # Attempt to get file extension from content type
                content_type_to_ext = {
                    'image/jpeg': '.jpg',
                    'image/png': '.png',
                    'image/gif': '.gif',
                    'image/webp': '.webp',
                    'image/svg+xml': '.svg',
                    'image/bmp': '.bmp',
                    'image/tiff': '.tiff'
                }
                
                # Get extension from URL or content type
                ext = os.path.splitext(img_url)[1].lower()
                if not ext or ext == '.':
                    ext = content_type_to_ext.get(content_type, '.jpg')
                
                # Create filename
                base_filename = os.path.basename(img_url)
                if not base_filename or base_filename == '':
                    # Create a unique name if URL doesn't provide one
                    import hashlib
                    hash_str = hashlib.md5(full_img_url.encode()).hexdigest()[:10]
                    base_filename = f"image_{hash_str}{ext}"
                elif not os.path.splitext(base_filename)[1]:
                    # Add extension if missing
                    base_filename += ext
                
                # Add prefix for flat structure
                if self.flat_structure:
                    filename = os.path.join(self.images_dir, f"img_{base_filename}")
                else:
                    filename = os.path.join(self.images_dir, base_filename)
                
                # Ensure the filename is safe and unique
                safe_filename = self._create_safe_filename(filename)
                
                # Write to file
                with open(safe_filename, 'wb') as f:
                    f.write(img_response.content)
                
                images_downloaded += 1
                self.logger.info(f"Successfully downloaded image: {safe_filename}")
                
                # Extract text from image
                self.extract_text_from_image(safe_filename)
                
            except requests.RequestException as e:
                self.logger.warning(f"Image download failed: {str(e)}")
                images_failed += 1
            except Exception as e:
                self.logger.warning(f"Image processing failed: {str(e)}")
                images_failed += 1
        
        self.logger.info(f"Images: {images_downloaded} downloaded, {images_skipped} skipped, {images_failed} failed")
        return images_downloaded
        
    def extract_text_from_image(self, filepath: str):
        """
        Extract text from image using OCR with robust error handling
        
        :param filepath: Path to the image file
        """
        # If image extraction is disabled, skip
        if not self.extract_images:
            return

        try:
            # Check if file exists and is accessible
            if not os.path.exists(filepath):
                self.logger.error(f"Image file not found: {filepath}")
                return
                
            # Check file size to avoid processing very large images
            file_size_mb = os.path.getsize(filepath) / (1024 * 1024)
            if file_size_mb > 10:  # 10MB limit
                self.logger.warning(f"Image file too large ({file_size_mb:.2f} MB), skipping OCR: {filepath}")
                return
            
            # Import pytesseract with error handling
            try:
                import pytesseract
                from PIL import Image, UnidentifiedImageError
            except ImportError as e:
                self.logger.error(f"Required library not found: {str(e)}. OCR processing skipped.")
                return
                
            try:
                # Open image with PIL
                image = Image.open(filepath)
                
                # Perform OCR
                text = pytesseract.image_to_string(image)
                
                # Clean text
                cleaned_text = self.clean_text(text)
                
                # Skip if no text was extracted
                if not cleaned_text or len(cleaned_text.strip()) < 5:
                    self.logger.info(f"No meaningful text extracted from image: {filepath}")
                    return
                
                # Add prefix for flat structure
                if self.flat_structure:
                    basename = f"img_ocr_{os.path.splitext(os.path.basename(filepath))[0]}.txt"
                else:
                    basename = f"{os.path.splitext(os.path.basename(filepath))[0]}.txt"
                    
                # Save text to file
                text_filename = os.path.join(self.images_dir, basename)
                with open(text_filename, 'w', encoding='utf-8') as f:
                    f.write(cleaned_text)
                    
                self.logger.info(f"Successfully extracted text from image: {filepath}")
                    
            except UnidentifiedImageError:
                self.logger.warning(f"Could not identify image format: {filepath}")
            except pytesseract.TesseractNotFoundError:
                self.logger.error("Tesseract OCR not installed or not in PATH")
                self.logger.error("Please install Tesseract OCR: https://github.com/tesseract-ocr/tesseract")
                # Disable image extraction to avoid repeated errors
                self.extract_images = False
            except Exception as e:
                self.logger.error(f"Error in OCR processing for {filepath}: {str(e)}")
                
        except Exception as e:
            self.logger.error(f"Error extracting text from image {filepath}: {str(e)}")
            
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text using spaCy if available, otherwise use basic regex
        Remove HTML tags, scripts, and unnecessary whitespace
        
        :param text: Input text to clean
        :return: Cleaned text
        """
        if not text:
            return ""
            
        # Remove HTML tags
        text = re.sub(r'<.*?>', '', text)
        
        # Process with spaCy if available
        if self.nlp is not None:
            try:
                doc = self.nlp(text)
                cleaned_text = ' '.join([token.text for token in doc if not token.is_space])
                return cleaned_text
            except Exception as e:
                self.logger.warning(f"spaCy processing failed: {str(e)}. Using basic cleaning.")
                # Fall back to basic cleaning
        
        # Basic cleaning if spaCy is not available or failed
        # Remove extra whitespace
        cleaned_text = re.sub(r'\s+', ' ', text)
        cleaned_text = cleaned_text.strip()
        
        return cleaned_text

    def save_webpage_text(self, soup: BeautifulSoup):
        """
        Save webpage text with strict extraction control
        """
        if not self.extract_text:
            return "Text not extracted."

        try:
            # Extract plain text
            text = soup.get_text(separator='\n', strip=True)
            
            # Optional: Clean text using spaCy
            doc = self.nlp(text)
            cleaned_text = ' '.join([token.text for token in doc if not token.is_space])
            
            # Add prefix for flat structure
            if self.flat_structure:
                text_file = os.path.join(self.output_dir, 'webpage_text_main.txt')
            else:
                text_file = os.path.join(self.output_dir, 'webpage_text.txt')
                
            # Save text
            with open(text_file, 'w', encoding='utf-8') as f:
                f.write(cleaned_text)
            
            self.logger.info("Webpage text extracted successfully")
            return "Text extracted."
        except Exception as e:
            self.logger.error(f"Text extraction failed: {e}")
            

    def scrape(self):
        """
        Comprehensive scraping method with strict selective extraction and error handling
        """
        self.logger.info(f"Starting scrape for URL: {self.url}")
        
        # Confirm extraction options
        self.logger.info(f"Extraction Options - Text: {self.extract_text}, "
                         f"Links: {self.extract_link}, "
                         f"Documents: {self.extract_documents}, "
                         f"Images: {self.extract_images}, "
                         f"Structure: {'Flat' if self.flat_structure else 'Nested'}")
        
        # Track overall success
        scrape_success = True
        summary = []
        
        # Fetch webpage
        try:
            response = self.fetch_page_content()
            
            if not response:
                self.logger.error("Scraping failed: Could not fetch webpage")
                return ["Failed to fetch webpage content"]
            
            # Parse HTML
            try:
                soup = BeautifulSoup(response.text, 'html.parser')
            except Exception as e:
                self.logger.error(f"Failed to parse HTML: {str(e)}")
                return ["Failed to parse HTML content"]
                
            # Selective extraction with error handling for each component
            if self.extract_link:
                try:
                    links = self.extract_links(soup)
                    saved_count = self.save_links(links)
                    summary.append(f"Links Extracted: {saved_count}")
                except Exception as e:
                    self.logger.error(f"Link extraction failed: {str(e)}")
                    summary.append("Link extraction failed")
                    scrape_success = False

            if self.extract_documents:
                try:
                    doc_count = self.download_documents(soup)
                    summary.append(f"Documents Downloaded: {doc_count}")
                except Exception as e:
                    self.logger.error(f"Document download failed: {str(e)}")
                    summary.append("Document download failed")
                    scrape_success = False
            
            if self.extract_images:
                try:
                    img_count = self.download_images(soup)
                    summary.append(f"Images Downloaded: {img_count}")
                except Exception as e:
                    self.logger.error(f"Image download failed: {str(e)}")
                    summary.append("Image download failed")
                    scrape_success = False
            
            if self.extract_text:
                try:
                    result = self.save_webpage_text(soup)
                    summary.append(result)
                except Exception as e:
                    self.logger.error(f"Text extraction failed: {str(e)}")
                    summary.append("Text extraction failed")
                    scrape_success = False
            
            # Log final status
            if scrape_success:
                self.logger.info("Scraping completed successfully")
                summary.append("Overall: Success")
            else:
                self.logger.warning("Scraping completed with some errors")
                summary.append("Overall: Partial success with errors")
            
            return summary
            
        except KeyboardInterrupt:
            self.logger.info("Scraping interrupted by user")
            return ["Scraping interrupted by user"]
            
        except Exception as e:
            self.logger.error(f"Unhandled exception during scraping: {str(e)}")
            import traceback
            self.logger.error(traceback.format_exc())
            return [f"Scraping failed: {str(e)}"]

    def _validate_url(self, url: str) -> bool:
        """
        Validate if a URL is safe to download from
        
        :param url: URL to validate
        :return: True if URL is valid and safe, False otherwise
        """
        try:
            # Basic validation
            parsed_url = urllib.parse.urlparse(url)
            
            # Check if scheme is http or https
            if parsed_url.scheme not in ('http', 'https'):
                self.logger.warning(f"Invalid URL scheme: {url}")
                return False
                
            # Check for IP address to avoid local network access
            import re
            import socket
            hostname = parsed_url.netloc.split(':')[0]
            is_ip = re.match(r"^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$", hostname)
            
            if is_ip:
                try:
                    ip_parts = [int(part) for part in hostname.split('.')]
                    # Check for private IP ranges
                    if (ip_parts[0] == 10 or  # 10.0.0.0/8
                        (ip_parts[0] == 172 and 16 <= ip_parts[1] <= 31) or  # 172.16.0.0/12
                        (ip_parts[0] == 192 and ip_parts[1] == 168) or  # 192.168.0.0/16
                        ip_parts[0] == 127):  # 127.0.0.0/8 (localhost)
                        self.logger.warning(f"Private IP address detected, skipping: {url}")
                        return False
                except:
                    pass
                    
            # Check content type for document and image downloads
            if self.extract_documents or self.extract_images:
                try:
                    # Send HEAD request to check content type
                    headers = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                    }
                    response = requests.head(url, headers=headers, timeout=5, allow_redirects=True)
                    
                    # Get content type
                    content_type = response.headers.get('Content-Type', '').lower()
                    
                    # Validate content type for safety
                    safe_types = [
                        'text/', 'image/', 'application/pdf', 'application/msword',
                        'application/vnd.openxmlformats-officedocument',
                        'application/vnd.ms-excel', 'application/vnd.ms-powerpoint',
                        'application/csv', 'text/csv'
                    ]
                    
                    if not any(safe_type in content_type for safe_type in safe_types):
                        self.logger.warning(f"Unsafe content type: {content_type} for URL: {url}")
                        return False
                        
                    # Check file size
                    content_length = response.headers.get('Content-Length')
                    if content_length and int(content_length) > 50 * 1024 * 1024:  # 50MB
                        self.logger.warning(f"File too large ({int(content_length) / (1024*1024):.2f} MB): {url}")
                        return False
                        
                except requests.RequestException:
                    # If HEAD request fails, we'll try GET when actually downloading
                    pass
                    
            return True
            
        except Exception as e:
            self.logger.warning(f"URL validation failed for {url}: {str(e)}")
            return False

def main():
    """
    Example usage for standalone script with robust error handling
    """
    import sys
    import argparse
    
    # Set up argument parser
    parser = argparse.ArgumentParser(description='Web Scraper Tool')
    parser.add_argument('--url', type=str, help='URL to scrape')
    parser.add_argument('--output', type=str, default=tempfile.gettempdir(), 
                        help='Output directory for scraped content')
    parser.add_argument('--text', action='store_true', help='Extract webpage text')
    parser.add_argument('--links', action='store_true', help='Extract hyperlinks')
    parser.add_argument('--docs', action='store_true', help='Download documents')
    parser.add_argument('--images', action='store_true', help='Download images')
    parser.add_argument('--flat', action='store_true', help='Use flat directory structure')
    parser.add_argument('--all', action='store_true', help='Enable all extraction options')
    
    # Parse arguments
    args = parser.parse_args()
    
    # Set default URL if not provided
    url = args.url or "https://cloud.google.com/learn/what-is-artificial-intelligence?hl=en"
    
    # Set extraction flags
    extract_all = args.all
    extract_text = args.text or extract_all
    extract_links = args.links or extract_all
    extract_docs = args.docs or extract_all
    extract_images = args.images or extract_all
    
    # If no extraction flags specified, enable all
    if not any([extract_text, extract_links, extract_docs, extract_images]):
        extract_text = extract_links = extract_docs = extract_images = True
    
    try:
        # Setup and execute scraper with nested structure (default)
        print(f"\nScraping URL: {url}")
        print(f"Output directory: {args.output}")
        print(f"Extraction flags - Text: {extract_text}, Links: {extract_links}, "
              f"Docs: {extract_docs}, Images: {extract_images}, "
              f"Structure: {'Flat' if args.flat else 'Nested'}")
        
        # Create scraper instance
        scraper = WebScraper(
            url,
            base_output_dir=args.output,
            extract_text=extract_text,
            extract_links=extract_links,
            extract_documents=extract_docs,
            extract_images=extract_images,
            flat_structure=args.flat
        )
        
        # Execute scraping
        results = scraper.scrape()
        
        # Print results
        print("\nScraping Results:")
        for item in results:
            print(f"- {item}")
        
        # Print output location
        print(f"\nScraped content saved to: {scraper.output_dir}")
        
    except ValueError as e:
        print(f"Error: {str(e)}")
        sys.exit(1)
    except KeyboardInterrupt:
        print("\nScraping interrupted by user")
        sys.exit(1)
    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        import traceback
        print(traceback.format_exc())
        sys.exit(1)

if __name__ == "__main__":
    main()